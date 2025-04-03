import os
import re

from docx import Document
from iotbx import mtz

from ....report.CCP4ReportParser import Report


class tableone_report(Report):

    TASKNAME = 'tableone'
    RUNNING = False


    def __init__(self, xmlnode=None, jobInfo={}, jobStatus=None, **kw):
        Report.__init__(self, xmlnode=xmlnode, jobInfo=jobInfo, jobStatus=jobStatus, **kw)
        if jobStatus is None or jobStatus.lower() == 'nooutput':
            return
        self.defaultReport()


    def GetPerformanceInfo(self):
        from ....core.CCP4ProjectsManager import PROJECTSMANAGER
        # Fetch the performance info related to the job that produced the input pdb file.
        dbm = PROJECTSMANAGER().db()
        for xdic in self.jobInfo["inputfiles"]:
            if xdic["filetypeclass"] == "PdbDataFile":
                pdbJobId = xdic["jobid"]
                break
        jobperf = dbm.getJobPerformance(pdbJobId)
        return jobperf


    def GetResolutionInfo(self):
        mtzpath = self.jobInfo["filenames"]["F_SIGF"]
        mobj = mtz.object(mtzpath)
        resol = mobj.max_min_resolution()
        return resol

    def getTable1L(self):
        namc = ["Resolution Range (Angstroms)", "Space Group", "Unit Cell (a, b, c), in Angstroms", "Unit Cell (alpha, beta, gamma), in degrees",
               "Total Reflections", "Unique Reflections", "Multiplicity",
               "Completeness (%)", "Signal Noise Ratio", "R-merge", "R-meas", "R-pim", "CC-1/2", "CC-anom"]
        datac = []
        resol = self.GetResolutionInfo()
        datac.append("%s to %s"%(str(round(resol[0],2)),str(round(resol[1], 2))))
        datac.append(self.xmlnode.findall('.//tableone/PdbInfo/SpaceGroup')[0].text)
        a = round(float(self.xmlnode.findall('.//tableone/PdbInfo/a')[0].text), 2)
        b = round(float(self.xmlnode.findall('.//tableone/PdbInfo/b')[0].text), 2)
        c = round(float(self.xmlnode.findall('.//tableone/PdbInfo/c')[0].text), 2)
        al = round(float(self.xmlnode.findall('.//tableone/PdbInfo/alpha')[0].text), 2)
        be = round(float(self.xmlnode.findall('.//tableone/PdbInfo/beta')[0].text), 2)
        ga = round(float(self.xmlnode.findall('.//tableone/PdbInfo/gamma')[0].text), 2)
        datac.append("(%s, %s, %s)"%(str(a), str(b), str(c)))
        datac.append("(%s, %s, %s)"%(str(al), str(be), str(ga)))
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Observ')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Unique')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Multi')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Compl')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/SigN')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Rmer')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Rmeas')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/Rpim')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/cchalf')[0].text)
        datac.append(self.xmlnode.findall('.//tableone/DataInfo/ccanom')[0].text)
        return namc, datac


    def getTable2L(self):
        datar = []
        namr = ["R-work", "R-free", "No. Atoms (exc. H)", "Protein Residues", "Rotamer Outliers (%)", "Molprobity Clashscore", "Mean B Factor",
                "Ramachandran Favoured Residues (%)", "Allowed Residues (%)", "Residues Not Favoured (%)"]
        jobperf = self.GetPerformanceInfo()
        datar.append(str(jobperf.get("RFactor")))
        datar.append(str(jobperf.get("RFree")))
        datar.append(self.xmlnode.findall('.//B_averages/Totals/Aminoacids/Atom_count')[0].text)
        datar.append(self.xmlnode.findall('.//Ramachandran_maps/Totals/Residues')[0].text)
        datar.append(self.xmlnode.findall('.//tableone/DataInfo/Rotamers')[0].text)
        datar.append(self.xmlnode.findall('.//tableone/DataInfo/Clashscore')[0].text)
        datar.append(str(round(float(self.xmlnode.findall('//B_averages/Totals/Aminoacids/Mean_B')[0].text), 3)))
        nRes = float(self.xmlnode.findall('.//Ramachandran_maps/Totals/Residues')[0].text)
        nfRes = float(self.xmlnode.findall('.//Ramachandran_maps/Totals/Favoured')[0].text)
        naRes = float(self.xmlnode.findall('.//Ramachandran_maps/Totals/Allowed')[0].text)
        datar.append(str(round(100.0*nfRes/nRes, 2)))
        datar.append(str(round(100.0*naRes/nRes, 2)))
        datar.append(str(round(100.0*(nRes-nfRes-naRes)/nRes, 2)))
        return namr, datar


    def defaultReport(self, parent=None):
        if parent is None:
            parent = self
        results = self.addResults()
        # Define Links & then tables.
        spath = os.path.join(self.jobInfo['fileroot'], "tabtest.docx")
        spathtex = os.path.join(self.jobInfo['fileroot'], "tabtest.txt")
        parent.append("<p>Successfully ran table one generation. Results summary below.</p>")
        parent.append('<span style="font-size:110%">Click on the following link to view the '
                                    'word document (which is located at : '
                                    + '<i>' + spath + '</i>' + ') </span>')
        parent.append('<br><a href="{0}"> &#8594;  Open docx file</a></br>'.format(spath))
        parent.append('<br><a href="{0}"> &#8594;  Open latex template</a></br>'.format(spathtex))
        parent.append('<br/>')
        # Construct Table lists pt1.
        namc, datac = self.getTable1L()
        DefFl = parent.addFold(label="Table One Information", initiallyOpen=True)
        SummaryTable = DefFl.addTable(select=".//tableone", style="width:260px;height:250px; float:left;")
        SummaryTable.addData(title="Data Collection Info", data=namc)
        SummaryTable.addData(title="Values", data=datac)
        # Construct Table lists pt2.
        namr, datar = self.getTable2L()
        refineTable = DefFl.addTable(select=".//tableone", style="width:260px;height:250px; float:left;")
        refineTable.addData(title="Refinement Statistics", data=namr)
        refineTable.addData(title="Values", data=datar)
        # Create Tables in required formats & save to spath/tex
        ccatlist = zip(namc+namr, datac+datar)
        self.CreateDocxTable(ccatlist, spath)
        self.CreateLatexTable(ccatlist, spathtex)

    def CreateDocxTable(self, ccatlist, spath):
        doc = Document()
        doc.add_heading('Table One', 0)
        table = doc.add_table(rows=len(ccatlist), cols=2)
        for i, tcont in enumerate(ccatlist):
            table.rows[i].cells[0].text = tcont[0]
            table.rows[i].cells[1].text = tcont[1]
        doc.save(spath)

    def CreateLatexTable(self, ccatlist, spath):
        txtio = open(spath, 'w')
        txtio.write("\\title{Tableone \\LaTeXe{} Template}\n")
        txtio.write("\\date{\\today}\n")
        txtio.write("\\documentclass[12pt]{article}\n")
        txtio.write("\\begin{document}\n")
        txtio.write("\\begin{table}[]\n")
        txtio.write("\\begin{tabular}{ll}\n")
        txtio.write("Meaurement & Value \\ \\hline\n")
        for tcont in ccatlist:
            rttl = re.sub(r"%","\%", tcont[0])
            txtio.write("%s & %s \\\\\n"%(rttl, tcont[1]))
        txtio.write("\\hline\n")
        txtio.write("\\end{tabular}\n")
        txtio.write("\\end{table}\n")
        txtio.write("\\end{document}\n")
        txtio.close()
